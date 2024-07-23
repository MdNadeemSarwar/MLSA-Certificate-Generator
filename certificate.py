import re
from docx.shared import Inches
import qrcode

# Function to replace text in document while preserving style
def replace_text(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            for run in p.runs:
                if regex.search(run.text):
                    run.text = regex.sub(replace, run.text)

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text(cell, regex, replace)

# Function to replace placeholder with participant name
def replace_participant_name(doc, name):
    regex = re.compile(r"PARTICIPANT NAME")
    replace_text(doc, regex, name)

# Function to replace placeholder with event name
def replace_event_name(doc, event):
    regex = re.compile(r"EVENT NAME")
    replace_text(doc, regex, event)

# Function to replace placeholder with host name
def replace_host_name(doc, name):
    regex = re.compile(r"HOST NAME")
    replace_text(doc, regex, name)

# Function to add QR code to the document
def add_qr_code(doc, qr_code_path):
    for paragraph in doc.paragraphs:
        if 'QR CODE' in paragraph.text:
            paragraph.clear()  # Clear the placeholder text
            run = paragraph.add_run()
            run.add_picture(qr_code_path, width=Inches(1.5))  # Add the QR code image
            break

# Function to generate QR code
def generate_qr_code(participant_name, event_name):
    qr_data = f"Participant: {participant_name}\nEvent: {event_name}"
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(qr_data)
    qr.make(fit=True)
    qr_img = qr.make_image(fill_color="black", back_color="white")
    qr_path = f'Output/QR/{participant_name}_qr.png'
    qr_img.save(qr_path)
    return qr_path
